from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application,
    CommandHandler,
    ConversationHandler,
    MessageHandler,
    filters,
    CallbackQueryHandler,
    ContextTypes,
)
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import sqlite3
import os
import re
from database import init_db, save_contract_data

# Этапы диалога
(
    CHOOSE_CONTRACT,
    INPUT_DATA,
    CHOOSE_FORMAT,
    END,
) = range(4)

# Словарь для хранения вопросов по каждому типу контракта
QUESTIONS = {
    "agent": [
        ("contract_name", "Введите название договора на русском языке:"),
        ("contract_name_en", "Введите название договора на английском языке:"),
        ("date", "Введите дату договора (ДД.ММ.ГГГГ):"),
        ("agent", "Введите название/ФИО Агента на русском языке:"),
        ("agent_en", "Введите название/ФИО Агента на английском языке:"),
        ("principal", "Введите название/ФИО Принципала на русском языке:"),
        ("principal_en", "Введите название/ФИО Принципала на английском языке:"),
        ("reward", "Введите размер вознаграждения (цифрами и прописью):"),
        ("reward_en", "Введите размер вознаграждения на английском:"),
        ("agent_details", "Введите реквизиты Агента:"),
        ("principal_details", "Введите реквизиты Принципала:"),
        ("poruchenie_subject", "Введите предмет поручения:"),
        ("poruchenie_deadline", "Введите сроки выполнения поручения:"),
        ("report_works", "Введите перечень выполненных работ для акта:"),
        ("report_amount", "Введите сумму отчета для акта:"),
    ],
    "subagent": [
        ("contract_name", "Введите название договора на русском языке:"),
        ("contract_name_en", "Введите название договора на английском языке:"),
        ("date", "Введите дату договора (ДД.ММ.ГГГГ):"),
        ("agent", "Введите название/ФИО Агента на русском языке:"),
        ("agent_en", "Введите название/ФИО Агента на английском языке:"),
        ("subagent", "Введите название/ФИО Субагента на русском языке:"),
        ("subagent_en", "Введите название/ФИО Субагента на английском языке:"),
        ("principal", "Введите название/ФИО Принципала на русском языке:"),
        ("principal_en", "Введите название/ФИО Принципала на английском языке:"),
        ("subject", "Введите предмет соглашения:"),
        ("reward", "Введите размер вознаграждения:"),
        ("agent_details", "Введите реквизиты Агента:"),
        ("principal_details", "Введите реквизиты Принципала:"),
        ("subagent_details", "Введите реквизиты Субагента:"),
        ("poruchenie_subject", "Введите предмет поручения:"),
        ("poruchenie_deadline", "Введите сроки выполнения поручения:"),
        ("report_works", "Введите перечень выполненных работ для акта:"),
    ],
    "supply": [
        ("contract_name", "Введите название договора на русском языке:"),
        ("contract_name_en", "Введите название договора на английском языке:"),
        ("date", "Введите дату договора (ДД.ММ.ГГГГ):"),
        ("supplier", "Введите название/ФИО Поставщика на русском языке:"),
        ("supplier_en", "Введите название/ФИО Поставщика на английском языке:"),
        ("purchaser", "Введите название/ФИО Покупателя на русском языке:"),
        ("purchaser_en", "Введите название/ФИО Покупателя на английском языке:"),
        ("supply_items", "Введите перечень товаров/услуг:"),
        ("price", "Введите цену и условия оплаты:"),
        ("delivery", "Введите сроки поставки:"),
        ("responsibility", "Введите ответственность сторон:"),
        ("agent_details", "Введите реквизиты Поставщика:"),
        ("principal_details", "Введите реквизиты Покупателя:"),
    ],
}


def check_english_fields(template_path):
    """Проверяет наличие английских полей в шаблоне."""
    if not os.path.exists(template_path):
        return False
    doc = Document(template_path)
    english_fields = [
        "{{NUMBER}}",
        "{{COMPANY}}",
        "{{AGENT}}",
        "{{PRINCIPAL}}",
        "{{SUBAGENT}}",
        "{{SUPPLIER}}",
        "{{PURCHASER}}",
        "{{REWARD}}",
        "{{SUBJECT}}",
    ]
    for paragraph in doc.paragraphs:
        if any(field in paragraph.text for field in english_fields):
            return True
    return False


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Начало диалога."""
    keyboard = [
        [InlineKeyboardButton("Агентское соглашение", callback_data="agent")],
        [InlineKeyboardButton("Субагентское соглашение", callback_data="subagent")],
        [InlineKeyboardButton("Договор поставки", callback_data="supply")],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Выберите тип договора:", reply_markup=reply_markup)
    return CHOOSE_CONTRACT


async def choose_contract(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Выбор типа договора."""
    query = update.callback_query
    await query.answer()
    contract_type = query.data
    context.user_data["contract_type"] = contract_type
    template_path = f"templates/{contract_type}_template.docx"
    context.user_data["template_path"] = template_path
    context.user_data["has_english"] = check_english_fields(template_path)
    context.user_data["question_index"] = 0
    context.user_data["answers"] = {}

    await ask_next_question(query.message, context)
    return INPUT_DATA


async def ask_next_question(message, context):
    """Задает следующий вопрос из списка."""
    user_data = context.user_data
    contract_type = user_data["contract_type"]
    question_index = user_data["question_index"]
    questions = QUESTIONS[contract_type]

    if question_index >= len(questions):
        await choose_format(message, context)
        return ConversationHandler.END

    field, question_text = questions[question_index]

    # Пропускаем английские вопросы, если не нужно
    if not user_data["has_english"] and field.endswith("_en"):
        user_data["question_index"] += 1
        await ask_next_question(message, context)
        return

    await message.reply_text(question_text)


async def handle_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обрабатывает ввод пользователя и задает следующий вопрос."""
    user_data = context.user_data
    contract_type = user_data["contract_type"]
    question_index = user_data["question_index"]
    questions = QUESTIONS[contract_type]
    field, _ = questions[question_index]
    answer = update.message.text

    # Валидация даты
    if field == "date" and not re.match(r"^\d{2}\.\d{2}\.\d{4}$", answer):
        await update.message.reply_text(
            "Неверный формат даты. Введите дату в формате ДД.ММ.ГГГГ:"
        )
        return INPUT_DATA

    user_data["answers"][field] = answer
    user_data["question_index"] += 1

    if user_data["question_index"] < len(questions):
        await ask_next_question(update.message, context)
        return INPUT_DATA
    else:
        await choose_format(update.message, context)
        return CHOOSE_FORMAT


async def choose_format(message, context):
    """Предлагает выбрать формат документа."""
    keyboard = [
        [
            InlineKeyboardButton("PDF", callback_data="pdf"),
            InlineKeyboardButton("DOCX", callback_data="docx"),
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await message.reply_text("Выберите формат для сохранения:", reply_markup=reply_markup)


async def generate_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Генерирует и отправляет документ."""
    query = update.callback_query
    await query.answer()
    file_format = query.data
    user_data = context.user_data
    user_id = query.from_user.id
    answers = user_data["answers"]
    contract_type = user_data["contract_type"]

    # Сохраняем данные в БД
    save_contract_data(user_id, {**user_data, **answers})

    # Генерация основного документа
    await generate_and_send_main_document(query.message, context, file_format)

    # Генерация доп. документов
    if contract_type in ["agent", "subagent"]:
        await generate_and_send_additional_documents(query.message, context)

    return ConversationHandler.END


async def generate_and_send_main_document(message, context, file_format):
    """Генерирует основной документ."""
    user_data = context.user_data
    answers = user_data["answers"]
    template_path = user_data["template_path"]
    doc = Document(template_path)

    # Замена плейсхолдеров
    for p in doc.paragraphs:
        for key, value in answers.items():
            placeholder = f"{{{{{key.upper()}}}}}"
            if placeholder in p.text:
                # Используем inline.text для сохранения форматирования
                for run in p.runs:
                    run.text = run.text.replace(placeholder, value)

    # Сохранение и отправка
    contract_name = answers.get("contract_name", "document")
    output_filename_base = f"output_{contract_name.replace(' ', '_')}"
    docx_path = f"{output_filename_base}.docx"
    doc.save(docx_path)

    if file_format == "pdf":
        pdf_path = f"{output_filename_base}.pdf"
        # Убедимся, что шрифт поддерживает кириллицу
        pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
        pdf = SimpleDocTemplate(pdf_path, pagesize=A4)
        styles = getSampleStyleSheet()
        styles['Normal'].fontName = 'DejaVuSans'
        content = [Paragraph(p.text, styles["Normal"]) for p in doc.paragraphs]
        pdf.build(content)
        with open(pdf_path, "rb") as f:
            await message.reply_document(f)
        os.remove(pdf_path)
    else:
        with open(docx_path, "rb") as f:
            await message.reply_document(f)

    os.remove(docx_path)


async def generate_and_send_additional_documents(message, context):
    """Генерирует и отправляет доп. документы (поручение, отчет)."""
    user_data = context.user_data
    answers = user_data["answers"]
    contract_type = user_data["contract_type"]
    contract_name = answers.get("contract_name", "document")

    for doc_type in ["assignment", "report"]:
        template_path = f"templates/{contract_type}_{doc_type}_template.docx"
        if not os.path.exists(template_path):
            continue

        doc = Document(template_path)
        for p in doc.paragraphs:
            for key, value in answers.items():
                placeholder = f"{{{{{key.upper()}}}}}"
                if placeholder in p.text:
                    for run in p.runs:
                        run.text = run.text.replace(placeholder, value)

        output_filename_base = f"output_{contract_name.replace(' ', '_')}_{doc_type}"
        docx_path = f"{output_filename_base}.docx"
        doc.save(docx_path)
        with open(docx_path, "rb") as f:
            await message.reply_document(f)
        os.remove(docx_path)


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Отменяет текущий диалог."""
    await update.message.reply_text("Действие отменено.")
    return ConversationHandler.END


def main():
    """Основная функция запуска бота."""
    # Убедитесь, что у вас есть токен бота
    token = os.getenv("TELEGRAM_TOKEN", "YOUR_TELEGRAM_BOT_TOKEN")
    if token == "YOUR_TELEGRAM_BOT_TOKEN":
        print("Ошибка: не найден токен Telegram. Установите переменную окружения TELEGRAM_TOKEN.")
        return

    # Инициализация БД
    init_db()

    # Установка шрифта для PDF
    # Убедитесь, что файл шрифта DejaVuSans.ttf находится в том же каталоге
    if not os.path.exists('DejaVuSans.ttf'):
        print("Ошибка: файл шрифта 'DejaVuSans.ttf' не найден. Скачайте его для корректной работы PDF.")
        # Попытка скачать шрифт
        try:
            import requests
            url = 'https://github.com/dejavu-fonts/dejavu-fonts/blob/main/ttf/DejaVuSans.ttf?raw=true'
            r = requests.get(url, allow_redirects=True)
            open('DejaVuSans.ttf', 'wb').write(r.content)
            print("Шрифт 'DejaVuSans.ttf' успешно скачан.")
        except Exception as e:
            print(f"Не удалось скачать шрифт: {e}")
            return


    app = Application.builder().token(token).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            CHOOSE_CONTRACT: [
                CallbackQueryHandler(choose_contract, pattern="^(agent|subagent|supply)$")
            ],
            INPUT_DATA: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_input)],
            CHOOSE_FORMAT: [
                CallbackQueryHandler(generate_document, pattern="^(pdf|docx)$")
            ],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
        map_to_parent={
            END: END,
        },
    )

    app.add_handler(conv_handler)
    app.run_polling()


if __name__ == "__main__":
    main()
