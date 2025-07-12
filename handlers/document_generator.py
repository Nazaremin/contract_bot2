import os
import shutil
from typing import Dict, List
from docx import Document
from datetime import datetime

from config import config
from db.database import db_manager

class DocumentGenerator:
    def __init__(self):
        self.templates_path = config.TEMPLATES_PATH
        self.output_path = config.OUTPUT_PATH
        os.makedirs(self.output_path, exist_ok=True)
    
    def _replace_placeholders(self, doc: Document, data: Dict) -> Document:
        """Замена плейсхолдеров в документе"""
        # Замена в параграфах
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if f"{{{key}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))
        
        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if f"{{{key}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{key}}}", str(value))
        
        return doc
    
    def _generate_single_document(self, template_name: str, data: Dict, output_name: str) -> str:
        """Генерация одного документа"""
        template_path = os.path.join(self.templates_path, template_name)
        output_path = os.path.join(self.output_path, output_name)
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Шаблон {template_name} не найден")
        
        # Создаем новый документ на основе шаблона
        doc = Document()
        
        # Читаем содержимое шаблона как текст и создаем параграфы
        with open(template_path, 'r', encoding='utf-8') as f:
            template_content = f.read()
        
        # Заменяем плейсхолдеры
        for key, value in data.items():
            template_content = template_content.replace(f"{{{key}}}", str(value))
        
        # Добавляем содержимое в документ
        for line in template_content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(output_path)
        return output_path
    
    def generate_agent_documents(self, data: Dict, user_id: int) -> List[str]:
        """Генерация документов для агентского соглашения"""
        contract_name = data.get('contract_name', 'contract')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        documents = []
        
        # Основной договор
        main_doc = self._generate_single_document(
            'agent_template.docx',
            data,
            f"agent_{contract_name}_{timestamp}.docx"
        )
        documents.append(main_doc)
        
        # Поручение
        assignment_doc = self._generate_single_document(
            'agent_assignment_template.docx',
            data,
            f"agent_assignment_{contract_name}_{timestamp}.docx"
        )
        documents.append(assignment_doc)
        
        # Акт отчета
        report_doc = self._generate_single_document(
            'agent_report_template.docx',
            data,
            f"agent_report_{contract_name}_{timestamp}.docx"
        )
        documents.append(report_doc)
        
        return documents
    
    def generate_subagent_documents(self, data: Dict, user_id: int) -> List[str]:
        """Генерация документов для субагентского соглашения"""
        contract_name = data.get('contract_name', 'contract')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        documents = []
        
        # Основной договор
        main_doc = self._generate_single_document(
            'subagent_template.docx',
            data,
            f"subagent_{contract_name}_{timestamp}.docx"
        )
        documents.append(main_doc)
        
        # Поручение
        assignment_doc = self._generate_single_document(
            'subagent_assignment_template.docx',
            data,
            f"subagent_assignment_{contract_name}_{timestamp}.docx"
        )
        documents.append(assignment_doc)
        
        # Акт отчета
        report_doc = self._generate_single_document(
            'subagent_report_template.docx',
            data,
            f"subagent_report_{contract_name}_{timestamp}.docx"
        )
        documents.append(report_doc)
        
        return documents
    
    def generate_delivery_documents(self, data: Dict, user_id: int) -> List[str]:
        """Генерация документов для договора поставки"""
        contract_name = data.get('contract_name', 'contract')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        documents = []
        
        # Основной договор
        main_doc = self._generate_single_document(
            'delivery_template.docx',
            data,
            f"delivery_{contract_name}_{timestamp}.docx"
        )
        documents.append(main_doc)
        
        return documents

document_generator = DocumentGenerator()

async def generate_documents(data: Dict, user_id: int) -> List[str]:
    """Основная функция генерации документов"""
    contract_type = data.get('contract_type')
    
    # Сохраняем данные в базу
    db_manager.save_contract(
        user_id=user_id,
        contract_type=contract_type,
        contract_name=data.get('contract_name', ''),
        data=data
    )
    
    # Генерируем документы
    if contract_type == 'agent':
        return document_generator.generate_agent_documents(data, user_id)
    elif contract_type == 'subagent':
        return document_generator.generate_subagent_documents(data, user_id)
    elif contract_type == 'delivery':
        return document_generator.generate_delivery_documents(data, user_id)
    else:
        raise ValueError(f"Неизвестный тип договора: {contract_type}")

